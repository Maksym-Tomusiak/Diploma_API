"""
Local Document Service - Handles operations on uploaded .docx files.
"""
from typing import Annotated, Optional
from dataclasses import dataclass, field
from io import BytesIO
from lxml import etree
import time
import zipfile

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml.simpletypes import ST_TwipsMeasure, ST_SignedTwipsMeasure
from docx.enum.text import WD_LINE_SPACING
from fastapi import Depends

from core.format_checker import FormatIssue, CheckResult
from core.document_formatter import FormatChange, FormatResult as FormatterResult
from schemas.template import TemplateParams


# Patch for parsing DOCX margins with float values
# 1. Save the original margin conversion methods
_orig_twips_convert = ST_TwipsMeasure.convert_from_xml
_orig_signed_twips_convert = ST_SignedTwipsMeasure.convert_from_xml


# 2. Create safe wrappers that clean the string if the original fails
@classmethod
def safe_twips_convert(cls, str_value):
    """
    Safely convert twips measurements from XML, handling float strings.
    Some DOCX files store margin values as float strings which causes parsing errors.
    """
    try:
        return _orig_twips_convert(str_value)
    except ValueError:
        # Cleans '1700.7874015748032' -> '1700' and passes it back to the original method
        return _orig_twips_convert(str(int(float(str_value))))


@classmethod
def safe_signed_twips_convert(cls, str_value):
    """
    Safely convert signed twips measurements from XML, handling float strings.
    """
    try:
        return _orig_signed_twips_convert(str_value)
    except ValueError:
        return _orig_signed_twips_convert(str(int(float(str_value))))


# 3. Apply the patch specifically to the margin measurement classes
ST_TwipsMeasure.convert_from_xml = safe_twips_convert
ST_SignedTwipsMeasure.convert_from_xml = safe_signed_twips_convert


@dataclass
class LocalTextSegment:
    """Represents a text segment from a local document."""
    content: str
    font_size_pt: Optional[float] = None
    font_family: Optional[str] = None
    char_count: int = 0
    paragraph_index: int = 0
    is_heading: bool = False
    is_on_first_page: bool = True


@dataclass
class ParagraphLineSpacing:
    """Represents line spacing for a paragraph with page information."""
    line_spacing: float
    paragraph_index: int
    is_on_first_page: bool = True


@dataclass
class LocalImageInfo:
    paragraph_index: int
    alignment: str  # 'center', 'left', 'right', 'justify', 'unknown'
    is_on_first_page: bool

@dataclass
class ParagraphAlignment:
    paragraph_index: int
    alignment: str
    text: str
    is_italic: bool
    is_on_first_page: bool

@dataclass
class LocalDocumentProperties:
    """Properties extracted from a local document."""
    title: str
    text_segments: list[LocalTextSegment] = field(default_factory=list)
    paragraph_line_spacings: list[ParagraphLineSpacing] = field(default_factory=list)
    margins: dict[str, float] = field(default_factory=dict)
    page_size: dict[str, float] = field(default_factory=dict)
    # Page numbering
    has_page_numbers: bool = False
    page_number_start: int = 1
    # First page different (skip first page for numbering/headers)
    first_page_different: bool = False
    # Image and alignment info
    images: list[LocalImageInfo] = field(default_factory=list)
    alignments: list[ParagraphAlignment] = field(default_factory=list)


class LocalDocumentService:
    """Service for working with local .docx files."""

    FONT_SIZE_TOLERANCE = 0.01
    MARGIN_TOLERANCE = 0.02  # inches (approx 0.5mm - allows for floating-point precision)
    LINE_SPACING_TOLERANCE = 0.01

    def _extract_font_from_rPr(self, rPr, theme_map: dict[str, str]) -> str | None:
        """Safely extracts the font name, prioritizing explicit declarations over stale themes."""
        if rPr is None:
            return None
            
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            return None

        # 1. PRIORITIZE EXPLICIT FONTS (Crucial for Google Docs & Cyrillic)
        # Check standard (ascii), High ANSI (hAnsi), and Complex Scripts (cs)
        explicit_font = (
            rFonts.get(qn('w:ascii')) or 
            rFonts.get(qn('w:hAnsi')) or 
            rFonts.get(qn('w:cs'))
        )
        
        if explicit_font:
            return explicit_font

        # 2. FALLBACK TO THEMES
        # Only use the theme if no explicit font was provided
        theme_key = (
            rFonts.get(qn('w:asciiTheme')) or 
            rFonts.get(qn('w:hAnsiTheme')) or 
            rFonts.get(qn('w:cstheme'))
        )
        
        if theme_key and theme_key in theme_map:
            return theme_map[theme_key]

        return None
    
    def _get_doc_defaults(self, doc: Document, theme_map: dict[str, str]) -> dict[str, str]:
        """
        Parses docDefaults and hunts down the 'Normal' style, 
        accounting for Google Docs omitting default flags entirely.
        """
        defaults_map = {}
        try:
            styles_element = doc.styles.element
            
            # Helper to search both direct rPr and paragraph-level rPr inside a style
            def extract_from_style_elem(elem):
                # 1. Check direct run properties
                rPr = elem.find(qn('w:rPr'))
                font = self._extract_font_from_rPr(rPr, theme_map)
                if font: return font
                
                # 2. Check paragraph properties -> run properties (Google Docs quirk)
                pPr = elem.find(qn('w:pPr'))
                if pPr is not None:
                    pPr_rPr = pPr.find(qn('w:rPr'))
                    return self._extract_font_from_rPr(pPr_rPr, theme_map)
                return None

            # 1. MICROSOFT WORD STANDARD (w:docDefaults)
            doc_defaults = styles_element.find(qn('w:docDefaults'))
            if doc_defaults is not None:
                rPrDefault = doc_defaults.find(qn('w:rPrDefault'))
                if rPrDefault is not None:
                    font = extract_from_style_elem(rPrDefault)
                    if font:
                        defaults_map['default_font'] = font
                        return defaults_map

            # 2. HUNT FOR THE NORMAL OR DEFAULT STYLE
            normal_font = None
            default_char_font = None
            
            for style in styles_element.findall(qn('w:style')):
                style_type = style.get(qn('w:type'))
                is_default = style.get(qn('w:default')) in ['1', 'true', 'on']
                style_id = style.get(qn('w:styleId'), "")
                
                name_elem = style.find(qn('w:name'))
                style_name = name_elem.get(qn('w:val'), "") if name_elem is not None else ""
                
                font = extract_from_style_elem(style)
                
                if font:
                    # If it explicitly says it's the default paragraph style, use it immediately
                    if is_default and style_type == 'paragraph':
                        defaults_map['default_font'] = font
                        return defaults_map
                        
                    # If it's a default character style, save it as a fallback
                    if is_default and style_type == 'character':
                        default_char_font = font
                        
                    # CRITICAL FIX: If it's named Normal, save it as the primary fallback
                    # This catches Google Docs exports that lack the w:default flag
                    if style_id.lower() == 'normal' or style_name.lower() == 'normal':
                        normal_font = font

            # 3. APPLY FALLBACKS (Normal > Default Character)
            if normal_font:
                defaults_map['default_font'] = normal_font
            elif default_char_font:
                defaults_map['default_font'] = default_char_font

        except Exception as e:
            print(f"Error parsing doc defaults: {e}")
            
        return defaults_map

    def _resolve_font_family(self, run, paragraph, theme_map, doc_defaults, expected_font_family: Optional[str] = None) -> str:
        """
        Resolves font family following strict inheritance hierarchy.
        """
        # 1. Direct Run Formatting
        try:
            if run._element.rPr is not None:
                font = self._extract_font_from_rPr(run._element.rPr, theme_map)
                if font: return font
        except AttributeError:
            pass

        # 2. Run (Character) Style Hierarchy
        if run.style:
            style = run.style
            while style:
                try:
                    if hasattr(style, '_element') and style._element is not None:
                        rPr = style._element.find(qn('w:rPr'))
                        font = self._extract_font_from_rPr(rPr, theme_map)
                        if font: return font
                except AttributeError:
                    pass
                style = getattr(style, 'base_style', None)

        # 3. Direct Paragraph Formatting 
        try:
            if paragraph._element.pPr is not None:
                pPr_rPr = paragraph._element.pPr.find(qn('w:rPr'))
                if pPr_rPr is not None:
                    font = self._extract_font_from_rPr(pPr_rPr, theme_map)
                    if font: return font
        except AttributeError:
            pass

        # 4. Paragraph Style Hierarchy
        if paragraph.style:
            style = paragraph.style
            while style:
                try:
                    if hasattr(style, '_element') and style._element is not None:
                        # A. Check direct style run properties
                        rPr = style._element.find(qn('w:rPr'))
                        font = self._extract_font_from_rPr(rPr, theme_map)
                        if font: return font
                        
                        # B. Check style paragraph properties (Google Docs style location)
                        pPr = style._element.find(qn('w:pPr'))
                        if pPr is not None:
                            pPr_rPr = pPr.find(qn('w:rPr'))
                            font = self._extract_font_from_rPr(pPr_rPr, theme_map)
                            if font: return font
                except AttributeError:
                    pass
                style = getattr(style, 'base_style', None)

        # 5. Document Defaults 
        # Because we now aggressively hunt for "Normal", this will catch the text 
        # even if python-docx lost track of the paragraph style entirely.
        if 'default_font' in doc_defaults:
            return doc_defaults['default_font']

        # 6. FINAL FALLBACK (ЗМІНЕНО ТУТ)
        # Якщо Google Docs взагалі не вказав шрифт, 
        # ми довіряємо expected_font_family (якщо він переданий), 
        # інакше беремо з теми документа.
        if expected_font_family:
            return expected_font_family
            
        return theme_map.get('minorHAnsi', 'Arial')
    
    def _resolve_line_spacing(self, paragraph) -> float:
        """Resolves line spacing, checking paragraph and style hierarchy."""
        if paragraph.paragraph_format.line_spacing is not None:
            spacing = paragraph.paragraph_format.line_spacing
            return float(spacing) if isinstance(spacing, (int, float)) else 1.0
        
        # Check style hierarchy
        style = paragraph.style
        while style:
            if style.paragraph_format.line_spacing is not None:
                spacing = style.paragraph_format.line_spacing
                return float(spacing) if isinstance(spacing, (int, float)) else 1.0
            style = getattr(style, 'base_style', None)
            
        return 1.15 # Standard Word/Docs default

    def _contains_page_number(self, section) -> bool:
        """Check if section headers/footers contain page number fields."""
        try:
            # Check header
            if section.header:
                header_xml = section.header._element
                # Look for fldChar with PAGE field
                for fld_char in header_xml.iter(qn('w:fldChar')):
                    return True
                # Look for instrText containing PAGE
                for instr_text in header_xml.iter(qn('w:instrText')):
                    if instr_text.text and 'PAGE' in instr_text.text:
                        return True
            
            # Check footer
            if section.footer:
                footer_xml = section.footer._element
                # Look for fldChar with PAGE field
                for fld_char in footer_xml.iter(qn('w:fldChar')):
                    return True
                # Look for instrText containing PAGE
                for instr_text in footer_xml.iter(qn('w:instrText')):
                    if instr_text.text and 'PAGE' in instr_text.text:
                        return True
        except Exception as e:
            print(f"Error checking page numbers: {e}")
        return False

    def _get_document_theme_fonts(self, doc: Document) -> dict[str, str]:
        """
        Parses the document's theme1.xml to find the actual font names,
        prioritizing Cyrillic fonts for Ukrainian/Russian text over standard Latin fonts.
        """
        theme_map = {
            'majorHAnsi': 'Calibri Light',
            'minorHAnsi': 'Calibri'
        }
        
        try:
            package_parts = doc.part.package.parts
            
            theme_part = None
            for part in package_parts:
                if part.content_type == 'application/vnd.openxmlformats-officedocument.theme+xml':
                    theme_part = part
                    break
            
            if theme_part:
                root = etree.fromstring(theme_part.blob)
                ns = {'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'}
                
                # Extract fonts for Major (Heading) and Minor (Body) schemes
                for scheme, key in [('majorFont', 'majorHAnsi'), ('minorFont', 'minorHAnsi')]:
                    scheme_elem = root.find(f'.//a:{scheme}', ns)
                    
                    if scheme_elem is not None:
                        # 1. Look for Cyrillic script explicitly first
                        cyrl_font = scheme_elem.find('./a:font[@script="Cyrl"]', ns)
                        if cyrl_font is not None and cyrl_font.get('typeface'):
                            theme_map[key] = cyrl_font.get('typeface')
                            continue  # Move to next scheme if we found Cyrillic
                            
                        # 2. Fallback to standard Latin if no Cyrillic is defined
                        latin_font = scheme_elem.find('./a:latin', ns)
                        if latin_font is not None and latin_font.get('typeface'):
                            theme_map[key] = latin_font.get('typeface')
                            
        except Exception as e:
            print(f"Warning: Could not parse document theme: {e}")
            
        return theme_map

    def extract_document_properties(self, file_content: bytes, expected_font_family: Optional[str] = None) -> LocalDocumentProperties:
        """
        Extract formatting properties from a .docx file.

        Args:
            file_content: Raw bytes of the .docx file

        Returns:
            LocalDocumentProperties with extracted formatting information
        """

        doc = Document(BytesIO(file_content))

        theme_fonts_map = self._get_document_theme_fonts(doc)

        doc_defaults = self._get_doc_defaults(doc, theme_fonts_map)

        # Get document title (from core properties or filename)
        title = doc.core_properties.title or "Untitled Document"

        text_segments = []
        paragraph_line_spacings = []
        images = []
        alignments = []

        # Get page dimensions for first page detection
        section = doc.sections[0] if doc.sections else None
        page_height_inches = section.page_height.inches if section and section.page_height else 11.0
        top_margin_inches = section.top_margin.inches if section and section.top_margin else 1.0
        bottom_margin_inches = section.bottom_margin.inches if section and section.bottom_margin else 1.0
        left_margin_inches = section.left_margin.inches if section and section.left_margin else 1.0
        right_margin_inches = section.right_margin.inches if section and section.right_margin else 1.0
        page_width_inches = section.page_width.inches if section and section.page_width else 8.5

        # Calculate usable dimensions for content on first page
        usable_page_height_inches = (page_height_inches - top_margin_inches - bottom_margin_inches) * 0.80
        usable_page_width_inches = page_width_inches - left_margin_inches - right_margin_inches

        # Estimate characters per line based on usable width
        # Using conservative estimate: ~8-9 chars per inch for typical 12pt fonts with spacing
        chars_per_line = max(40, int(usable_page_width_inches * 8.5))

        # --- NEW: Track explicit page boundaries ---
        cumulative_height_inches = 0.0
        has_passed_first_page = False

        # Extract text and formatting from paragraphs
        for para_idx, paragraph in enumerate(doc.paragraphs):
            # Check if paragraph is a heading
            is_heading = paragraph.style.name.startswith('Heading')

            # Estimate paragraph height (very rough approximation)
            # Average font size of runs in paragraph, or default to 12pt
            para_font_sizes = [run.font.size.pt for run in paragraph.runs if run.font.size]
            avg_font_size = sum(para_font_sizes) / len(para_font_sizes) if para_font_sizes else 12.0

            # Line height is approximately font size * line spacing (default 1.15)
            line_spacing_multiple = 1.15
            if paragraph.paragraph_format.line_spacing is not None:
                spacing = paragraph.paragraph_format.line_spacing
                if isinstance(spacing, (int, float)):
                    line_spacing_multiple = float(spacing)

            # Calculate line height in inches (1 pt = 1/72 inch)
            line_height_inches = (avg_font_size / 72.0) * line_spacing_multiple

            # Count approximate lines in paragraph based on actual page width
            para_text_length = len(paragraph.text)
            estimated_lines = max(1, para_text_length // chars_per_line) if para_text_length > 0 else 1

            # Add paragraph height
            para_height_inches = line_height_inches * estimated_lines

            # Add space before/after if specified
            if paragraph.paragraph_format.space_before:
                para_height_inches += paragraph.paragraph_format.space_before.inches
            if paragraph.paragraph_format.space_after:
                para_height_inches += paragraph.paragraph_format.space_after.inches

            # --- NEW FIRST PAGE DETECTION LOGIC ---
            # 1. Check for explicit page boundaries before this paragraph
            if not has_passed_first_page:
                if paragraph.paragraph_format.page_break_before:
                    has_passed_first_page = True
                else:
                    for run in paragraph.runs:
                        # Catch pagination boundaries established by Word/Docs
                        if run._element.find(qn('w:lastRenderedPageBreak')) is not None:
                            has_passed_first_page = True
                            break

            # 2. Fallback to height heuristic if no hard break is found
            if not has_passed_first_page:
                if cumulative_height_inches > usable_page_height_inches:
                    has_passed_first_page = True

            is_on_first_page = not has_passed_first_page
            cumulative_height_inches += para_height_inches

            # 3. Check for manual page breaks INSIDE this paragraph (affects the NEXT paragraph)
            if is_on_first_page:
                for run in paragraph.runs:
                    for br in run._element.findall(qn('w:br')):
                        if br.get(qn('w:type')) == 'page':
                            has_passed_first_page = True
                            break
            # ----------------------------------------

            # Track line spacing for paragraph (after is_on_first_page is determined)
            spacing = self._resolve_line_spacing(paragraph)
            paragraph_line_spacings.append(ParagraphLineSpacing(
                line_spacing=spacing,
                paragraph_index=para_idx,
                is_on_first_page=is_on_first_page,
            ))

            # --- NEW: Image Detection ---
            has_image = False
            # Check for drawings in paragraph XML
            if paragraph._element.find(qn('w:drawing')) is not None or \
               paragraph._element.find(qn('w:pict')) is not None:
                has_image = True
            
            # Also check runs for drawings (more reliable for inline images)
            if not has_image:
                for run in paragraph.runs:
                    if run._element.find(qn('w:drawing')) is not None or \
                       run._element.find(qn('w:pict')) is not None:
                        has_image = True
                        break
            
            # Resolve paragraph alignment
            alignment = "unknown"
            if paragraph.alignment is not None:
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                if paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER: alignment = "center"
                elif paragraph.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY: alignment = "justify"
                elif paragraph.alignment == WD_ALIGN_PARAGRAPH.LEFT: alignment = "left"
                elif paragraph.alignment == WD_ALIGN_PARAGRAPH.RIGHT: alignment = "right"
            elif paragraph.style and hasattr(paragraph.style, 'paragraph_format'):
                # Check style alignment
                style_align = paragraph.style.paragraph_format.alignment
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                if style_align == WD_ALIGN_PARAGRAPH.CENTER: alignment = "center"
                elif style_align == WD_ALIGN_PARAGRAPH.JUSTIFY: alignment = "justify"
                elif style_align == WD_ALIGN_PARAGRAPH.LEFT: alignment = "left"
                elif style_align == WD_ALIGN_PARAGRAPH.RIGHT: alignment = "right"

            if has_image:
                images.append(LocalImageInfo(
                    paragraph_index=para_idx,
                    alignment=alignment,
                    is_on_first_page=is_on_first_page
                ))
            
            # Store alignment info for all paragraphs to check captions/sources
            is_italic = any(run.font.italic for run in paragraph.runs if run.text.strip())
            alignments.append(ParagraphAlignment(
                paragraph_index=para_idx,
                alignment=alignment,
                text=paragraph.text.strip(),
                is_italic=is_italic,
                is_on_first_page=is_on_first_page
            ))

            # Extract runs (text segments with consistent formatting)
            for run in paragraph.runs:
                if not run.text.strip():
                    continue

                # Get font size
                font_size = None
                if run.font.size:
                    font_size = run.font.size.pt

                font_family = self._resolve_font_family(run, paragraph, theme_fonts_map, doc_defaults, expected_font_family)

                segment = LocalTextSegment(
                    content=run.text,
                    font_size_pt=font_size,
                    font_family=font_family,
                    char_count=len(run.text),
                    paragraph_index=para_idx,
                    is_heading=is_heading,
                    is_on_first_page=is_on_first_page,
                )
                text_segments.append(segment)

        # Extract margins (in inches, convert from twips)
        margins = {}
        section = doc.sections[0] if doc.sections else None
        if section:
            margins = {
                "top": section.top_margin.inches if section.top_margin else 1.0,
                "bottom": section.bottom_margin.inches if section.bottom_margin else 1.0,
                "left": section.left_margin.inches if section.left_margin else 1.0,
                "right": section.right_margin.inches if section.right_margin else 1.0,
            }

            # Page size
            page_size = {
                "width": section.page_width.inches if section.page_width else 8.5,
                "height": section.page_height.inches if section.page_height else 11.0,
            }
        else:
            page_size = {"width": 8.5, "height": 11.0}

        # Check for page numbers in headers/footers
        has_page_numbers = False
        page_number_start = 1
        first_page_different = False
        
        if section:
            has_page_numbers = self._contains_page_number(section)
            
            # Check if first page is different
            try:
                sectPr = section._sectPr
                titlePg = sectPr.find(qn('w:titlePg'))
                first_page_different = titlePg is not None
            except Exception:
                pass
            
            # Try to get page number start (pgNumType)
            try:
                sectPr = section._sectPr
                pgNumType = sectPr.find(qn('w:pgNumType'))
                if pgNumType is not None:
                    start_attr = pgNumType.get(qn('w:start'))
                    if start_attr:
                        page_number_start = int(start_attr)
            except Exception:
                pass

        return LocalDocumentProperties(
            title=title,
            text_segments=text_segments,
            paragraph_line_spacings=paragraph_line_spacings,
            margins=margins,
            page_size=page_size,
            has_page_numbers=has_page_numbers,
            page_number_start=page_number_start,
            first_page_different=first_page_different,
            images=images,
            alignments=alignments
        )

    def debug_google_docs_xml(self, file_content: bytes):
        doc = Document(BytesIO(file_content))
        styles_element = doc.styles.element
        
        # Helper to safely print the raw XML
        def print_xml(elem):
            if elem is not None:
                print(etree.tostring(elem, pretty_print=True, encoding='unicode'))
            else:
                print("Element is None")

        print("=== 1. DOC DEFAULTS ===")
        doc_defaults = styles_element.find(qn('w:docDefaults'))
        if doc_defaults is not None:
            print_xml(doc_defaults)
        else:
            print("No w:docDefaults found.")
            
        print("\n=== 2. RELEVANT STYLES ===")
        for style in styles_element.findall(qn('w:style')):
            style_id = style.get(qn('w:styleId'))
            name_node = style.find(qn('w:name'))
            name = name_node.get(qn('w:val')) if name_node is not None else ""
            is_default = style.get(qn('w:default'))
            
            # Grab anything marked default, or named Normal
            if is_default in ['1', 'true'] or style_id == 'Normal' or name == 'Normal':
                print(f"\n--- Style: ID='{style_id}', Name='{name}', Default='{is_default}' ---")
                print_xml(style)

        print("\n=== 3. FIRST PARAGRAPH XML ===")
        # Find the first paragraph that actually contains text to see its properties
        for p in doc.paragraphs:
            if p.text.strip():
                print_xml(p._element)
                break

    def check_document(
        self,
        file_content: bytes,
        params: TemplateParams,
        expected_font_family: Optional[str] = None,
    ) -> CheckResult:
        """
        Check a local document's formatting against parameters.

        Args:
            file_content: Raw bytes of the .docx file
            params: Expected formatting parameters
            expected_font_family: Expected font family name

        Returns:
            CheckResult with formatting issues found
        """
        start_time = time.time()
        issues: list[FormatIssue] = []

        #self.debug_google_docs_xml(file_content)

        doc_props = self.extract_document_properties(file_content, expected_font_family)

        # Font Size & Family Checks
        font_size_issues: list[LocalTextSegment] = []
        font_family_issues: list[LocalTextSegment] = []

        for segment in doc_props.text_segments:
            # Skip first page content if requested
            if params.skip_first_page and segment.is_on_first_page:
                print(segment.content)
                continue

            if segment.font_size_pt is not None:
                size_diff = abs(segment.font_size_pt - params.font_size)
                if size_diff > self.FONT_SIZE_TOLERANCE:
                    font_size_issues.append(segment)

            if expected_font_family and segment.font_family:
                # Only check segments that have an explicitly set font
                if segment.font_family.lower() != expected_font_family.lower():
                    font_family_issues.append(segment)

        # Report Font Size Issues
        if font_size_issues:
            size_groups: dict[float, list[LocalTextSegment]] = {}
            for seg in font_size_issues:
                size = seg.font_size_pt or 0
                if size not in size_groups:
                    size_groups[size] = []
                size_groups[size].append(seg)

            for actual_size, segments in size_groups.items():
                total_chars_this_size = sum(s.char_count for s in segments)
                excerpts = [f'"{s.content[:50]}..."' for s in segments[:3]]
                excerpt_text = ", ".join(excerpts)

                issues.append(FormatIssue(
                    type="font_size_mismatch",
                    severity="high" if abs(actual_size - params.font_size) > 2 else "medium",
                    details=f"Found {total_chars_this_size} characters with font size {actual_size:.1f}pt (expected {params.font_size}pt). Examples: {excerpt_text}",
                    expected=f"{params.font_size}pt",
                    actual=f"{actual_size:.1f}pt",
                ))

        # Report Font Family Issues
        if font_family_issues:
            font_groups: dict[str, list[LocalTextSegment]] = {}
            for seg in font_family_issues:
                font = seg.font_family or "Unknown"
                if font not in font_groups:
                    font_groups[font] = []
                font_groups[font].append(seg)

            for actual_font, segments in font_groups.items():
                total_chars_this_font = sum(s.char_count for s in segments)
                excerpts = [f'"{s.content[:50]}..."' for s in segments[:3]]
                excerpt_text = ", ".join(excerpts)

                issues.append(FormatIssue(
                    type="font_family_mismatch",
                    severity="high",
                    details=f"Found {total_chars_this_font} characters with font '{actual_font}' (expected '{expected_font_family}'). Examples: {excerpt_text}",
                    expected=expected_font_family,
                    actual=actual_font,
                ))

        # Line Spacing Checks
        if doc_props.paragraph_line_spacings:
            wrong_spacing_values: dict[float, int] = {}
            for para_spacing in doc_props.paragraph_line_spacings:
                # Skip first page content if requested
                if params.skip_first_page and para_spacing.is_on_first_page:
                    continue
                
                spacing_diff = abs(para_spacing.line_spacing - params.line_spacing)
                if spacing_diff > self.LINE_SPACING_TOLERANCE:
                    spacing_value = para_spacing.line_spacing
                    wrong_spacing_values[spacing_value] = wrong_spacing_values.get(spacing_value, 0) + 1

            for actual_spacing, count in wrong_spacing_values.items():
                issues.append(FormatIssue(
                    type="line_spacing_mismatch",
                    severity="medium",
                    details=f"{count} paragraphs have line spacing {actual_spacing:.2f} (expected {params.line_spacing:.2f})",
                    expected=f"{params.line_spacing:.2f}",
                    actual=f"{actual_spacing:.2f}",
                ))

        # Margin Checks
        if doc_props.margins and params.margins:
            # Convert mm to inches for comparison (params.margins are in mm, doc_props.margins are in inches)
            def mm_to_inches(mm: float) -> float:
                return mm / 25.4
            
            margin_mapping = {
                'top': params.margins.top,
                'bottom': params.margins.bottom,
                'left': params.margins.left,
                'right': params.margins.right,
            }
            
            for margin_name, expected_mm in margin_mapping.items():
                if margin_name in doc_props.margins:
                    actual_inches = doc_props.margins[margin_name]
                    expected_inches = mm_to_inches(expected_mm)
                    diff = abs(actual_inches - expected_inches)
                    if diff > self.MARGIN_TOLERANCE:
                        actual_mm = actual_inches * 25.4
                        issues.append(FormatIssue(
                            type=f"margin_{margin_name}_mismatch",
                            severity="medium",
                            details=f"{margin_name.capitalize()} margin is {actual_mm:.1f}mm (expected {expected_mm:.1f}mm)",
                            expected=f"{expected_mm:.1f}mm",
                            actual=f"{actual_mm:.1f}mm",
                        ))

        # Image and Caption Checks
        import re
        # Strict pattern for Рис. X.X. (exactly two numbers)
        caption_pattern = re.compile(r'^(Рис\.|Зоб\.|Фото|Рисунок)\s+\d+\.\d+\.?\s+.+')
        
        for img in doc_props.images:
            if params.skip_first_page and img.is_on_first_page:
                continue
                
            # 1. Image alignment
            if img.alignment != 'center':
                # Try to get caption for better identification
                caption_para = next((a for a in doc_props.alignments if a.paragraph_index == img.paragraph_index + 1), None)
                img_ref = "Зображення"
                if caption_para and caption_para.text:
                    # Extract the first part like "Рис. 1.1."
                    match = re.match(r'^(Рис\.|Зоб\.|Рисунок|Фото)\s+\d+\.\d+\.?\s*', caption_para.text)
                    if match:
                        img_ref = match.group(0).strip()
                    else:
                        img_ref = f"Зображення '{caption_para.text[:20]}...'"

                issues.append(FormatIssue(
                    type="image_alignment_error",
                    severity="medium",
                    details=f"{img_ref} має бути вирівняно по центру",
                    expected="center",
                    actual=img.alignment
                ))
            
            # 2. Check for caption (next paragraph)
            caption_para = next((a for a in doc_props.alignments if a.paragraph_index == img.paragraph_index + 1), None)
            if not caption_para or not caption_para.text:
                issues.append(FormatIssue(
                    type="image_caption_missing",
                    severity="high",
                    details=f"Відсутній підпис під зображенням у параграфі {img.paragraph_index + 2}",
                    expected="Рис. X.X. Назва",
                    actual="порожньо"
                ))
            else:
                # Check format
                if not caption_pattern.match(caption_para.text):
                    issues.append(FormatIssue(
                        type="image_caption_format_error",
                        severity="low",
                        details=f"Неправильний формат підпису: '{caption_para.text[:30]}...'. Очікується: 'Рис. X.X. Опис'",
                        expected="Рис. X.X. Назва",
                        actual=caption_para.text[:30]
                    ))
                
                # Check alignment
                if caption_para.alignment != "center":
                    issues.append(FormatIssue(
                        type="image_caption_alignment_error",
                        severity="low",
                        details=f"Підпис під зображенням '{caption_para.text[:20]}' має бути відцентрований",
                        expected="center",
                        actual=caption_para.alignment
                    ))

                # 3. Check for source (paragraph after caption)
                source_para = next((a for a in doc_props.alignments if a.paragraph_index == img.paragraph_index + 2), None)
                if not source_para or not source_para.text:
                    issues.append(FormatIssue(
                        type="image_source_missing",
                        severity="medium",
                        details=f"Відсутнє джерело під зображенням '{caption_para.text[:20]}...'",
                        expected="Джерело: розроблено автором (або інше)",
                        actual="порожньо"
                    ))
                elif not source_para.text.lower().startswith("джерело:"):
                    issues.append(FormatIssue(
                        type="image_source_format_error",
                        severity="low",
                        details=f"Неправильний формат джерела: '{source_para.text[:30]}...'. Очікується: 'Джерело: опис'",
                        expected="Джерело: ...",
                        actual=source_para.text[:30]
                    ))
                else:
                    # If it exists and starts with "Джерело:", check alignment and style
                    if source_para.alignment != "center":
                        issues.append(FormatIssue(
                            type="image_source_alignment_error",
                            severity="low",
                            details=f"Рядок джерела '{source_para.text[:20]}' має бути відцентрований",
                            expected="center",
                            actual=source_para.alignment
                        ))
                    if not source_para.is_italic:
                        issues.append(FormatIssue(
                            type="image_source_style_error",
                            severity="low",
                            details=f"Рядок джерела '{source_para.text[:20]}' має бути написаний курсивом",
                            expected="italic",
                            actual="normal"
                        ))

        # Page Numbering Checks
        if params.check_numbering:
            if not doc_props.has_page_numbers:
                issues.append(FormatIssue(
                    type="page_numbering_missing",
                    severity="high",
                    details="Document does not have page numbers, but they are required",
                    expected="Page numbers enabled",
                    actual="No page numbers found",
                ))
            else:
                # Check page number start
                expected_start = params.start_from_number
                actual_start = doc_props.page_number_start
                
                if params.skip_first_page:
                    # When skip_first_page is true, first page shouldn't have numbering
                    # Page 2 should show the expected start number
                    if not doc_props.first_page_different:
                        issues.append(FormatIssue(
                            type="page_numbering_on_first_page",
                            severity="medium",
                            details="First page numbering should be hidden. Enable 'Different first page' in Page Setup.",
                            expected="First page without number",
                            actual="First page has number",
                        ))
                    elif actual_start != expected_start:
                        issues.append(FormatIssue(
                            type="page_number_start_mismatch",
                            severity="medium",
                            details=f"Page numbering starts at {actual_start}, expected {expected_start}. The number shown on page 2 should be {expected_start}.",
                            expected=f"{expected_start}",
                            actual=f"{actual_start}",
                        ))
                else:
                    # When skip_first_page is false, numbering should start on first page
                    if doc_props.first_page_different:
                        issues.append(FormatIssue(
                            type="page_numbering_first_page_different",
                            severity="medium",
                            details="First page is set to be different ('Different first page' is enabled). This hides page numbering on the first page. Please disable 'Different first page' to ensure numbering starts on Page 1.",
                            expected="Same header/footer on all pages",
                            actual="Different first page enabled",
                        ))
                    elif actual_start != expected_start:
                        issues.append(FormatIssue(
                            type="page_number_start_mismatch",
                            severity="medium",
                            details=f"Page numbering starts at {actual_start}, expected {expected_start}",
                            expected=f"{expected_start}",
                            actual=f"{actual_start}",
                        ))

        # Calculate score
        if not doc_props.text_segments:
            issues.append(FormatIssue(
                type="no_text_found",
                severity="low",
                details="No text content found in document",
            ))

        # Score calculation
        critical_issues = len([i for i in issues if i.severity == "high"])
        medium_issues = len([i for i in issues if i.severity == "medium"])
        low_issues = len([i for i in issues if i.severity == "low"])

        total_checks = max(1, critical_issues + medium_issues + low_issues + 10)
        penalty = (critical_issues * 3) + (medium_issues * 1.5) + (low_issues * 0.5)
        overall_score = max(0.0, (total_checks - penalty) / total_checks)  # Return as fraction 0.0-1.0

        processing_time = int((time.time() - start_time) * 1000)

        return CheckResult(
            passed=len(issues) == 0,
            overall_score=overall_score,
            issues=issues,
            processing_time_ms=processing_time,
            document_title=doc_props.title,
        )
    def format_document(
        self,
        file_content: bytes,
        params: TemplateParams,
        expected_font_family: Optional[str] = None,
    ) -> tuple[bytes, FormatterResult]:
        """
        Apply formatting to a local document and return the modified file.
        
        Args:
            file_content: Raw bytes of the .docx file
            params: Formatting parameters to apply
            expected_font_family: Font family to apply
            
        Returns:
            Tuple of (modified_file_bytes, FormatterResult)
        """
        start_time = time.time()
        changes: list[FormatChange] = []
        
        doc = Document(BytesIO(file_content))
        doc_title = doc.core_properties.title or "Untitled Document"

        theme_fonts_map = self._get_document_theme_fonts(doc)
        doc_defaults = self._get_doc_defaults(doc, theme_fonts_map)
        
        # Get page dimensions for first page detection (if skip_first_page is enabled)
        section = doc.sections[0] if doc.sections else None
        page_height_inches = section.page_height.inches if section and section.page_height else 11.0
        top_margin_inches = section.top_margin.inches if section and section.top_margin else 1.0
        bottom_margin_inches = section.bottom_margin.inches if section and section.bottom_margin else 1.0
        left_margin_inches = section.left_margin.inches if section and section.left_margin else 1.0
        right_margin_inches = section.right_margin.inches if section and section.right_margin else 1.0
        page_width_inches = section.page_width.inches if section and section.page_width else 8.5
        
        # Calculate usable dimensions for first page detection
        usable_page_height_inches = (page_height_inches - top_margin_inches - bottom_margin_inches) * 0.92
        usable_page_width_inches = page_width_inches - left_margin_inches - right_margin_inches
        chars_per_line = max(40, int(usable_page_width_inches * 8.5))
        
        # --- NEW: Track explicit page boundaries ---
        cumulative_height_inches = 0.0
        has_passed_first_page = False
        
        # Apply formatting to all paragraphs
        for para_idx, paragraph in enumerate(doc.paragraphs):
            # Check if paragraph is a heading
            is_heading = paragraph.style.name.startswith('Heading')
            
            # Estimate paragraph height for first page detection
            para_font_sizes = [run.font.size.pt for run in paragraph.runs if run.font.size]
            avg_font_size = sum(para_font_sizes) / len(para_font_sizes) if para_font_sizes else 12.0
            
            line_spacing_multiple = 1.15
            if paragraph.paragraph_format.line_spacing is not None:
                spacing = paragraph.paragraph_format.line_spacing
                if isinstance(spacing, (int, float)):
                    line_spacing_multiple = float(spacing)
            
            line_height_inches = (avg_font_size / 72.0) * line_spacing_multiple
            para_text_length = len(paragraph.text)
            estimated_lines = max(1, para_text_length // chars_per_line) if para_text_length > 0 else 1
            para_height_inches = line_height_inches * estimated_lines
            
            if paragraph.paragraph_format.space_before:
                para_height_inches += paragraph.paragraph_format.space_before.inches
            if paragraph.paragraph_format.space_after:
                para_height_inches += paragraph.paragraph_format.space_after.inches
            
            # --- NEW FIRST PAGE DETECTION LOGIC ---
            # 1. Check for explicit page boundaries before this paragraph
            if not has_passed_first_page:
                if paragraph.paragraph_format.page_break_before:
                    has_passed_first_page = True
                else:
                    for run in paragraph.runs:
                        # Catch pagination boundaries established by Word/Docs
                        if run._element.find(qn('w:lastRenderedPageBreak')) is not None:
                            has_passed_first_page = True
                            break

            # 2. Fallback to height heuristic if no hard break is found
            if not has_passed_first_page:
                if cumulative_height_inches > usable_page_height_inches:
                    has_passed_first_page = True

            is_on_first_page = not has_passed_first_page
            cumulative_height_inches += para_height_inches

            # 3. Check for manual page breaks INSIDE this paragraph (affects the NEXT paragraph)
            if is_on_first_page:
                for run in paragraph.runs:
                    for br in run._element.findall(qn('w:br')):
                        if br.get(qn('w:type')) == 'page':
                            has_passed_first_page = True
                            break
            # ----------------------------------------
            
            # Skip first page content if requested
            if params.skip_first_page and is_on_first_page:
                continue
            
            # Apply line spacing
            if paragraph.paragraph_format.line_spacing != params.line_spacing:
                old_spacing = paragraph.paragraph_format.line_spacing
                paragraph.paragraph_format.line_spacing = params.line_spacing
                changes.append(FormatChange(
                    type="line_spacing",
                    description=f"Updated line spacing for paragraph",
                    before=f"{old_spacing:.2f}" if old_spacing else "unset",
                    after=f"{params.line_spacing:.2f}",
                ))
            
            # Apply formatting to runs
            for run in paragraph.runs:
                if not run.text.strip():
                    continue
                
                # Ensure headings are explicitly bold
                if is_heading and not run.font.bold:
                    run.font.bold = True
                
                # Apply font size
                if run.font.size:
                    old_size = run.font.size.pt
                    if abs(old_size - params.font_size) > self.FONT_SIZE_TOLERANCE:
                        run.font.size = Pt(params.font_size)
                        changes.append(FormatChange(
                            type="font_size",
                            description=f"Changed font size",
                            before=f"{old_size:.1f}pt",
                            after=f"{params.font_size}pt",
                        ))
                else:
                    run.font.size = Pt(params.font_size)
                    changes.append(FormatChange(
                        type="font_size",
                        description=f"Set font size",
                        before="unset",
                        after=f"{params.font_size}pt",
                    ))
                
                # Apply font family
                if expected_font_family:
                    # 1. Use the robust resolver to check the REAL current font
                    old_font = self._resolve_font_family(run, paragraph, theme_fonts_map, doc_defaults, expected_font_family)
                    
                    if old_font.lower() != expected_font_family.lower():
                        # 2. Set the python-docx property (Updates w:ascii mostly)
                        run.font.name = expected_font_family
                        
                        # 3. CRITICAL: Manipulate XML to remove Theme References
                        # Access the raw XML element for the run's properties
                        r = run._element
                        rPr = r.get_or_add_rPr()
                        rFonts = rPr.get_or_add_rFonts()
                        
                        # A. Force the font on ALL slots (East Asia often overrides)
                        rFonts.set(qn('w:ascii'), expected_font_family)
                        rFonts.set(qn('w:hAnsi'), expected_font_family)
                        rFonts.set(qn('w:eastAsia'), expected_font_family)
                        rFonts.set(qn('w:cs'), expected_font_family)
                        
                        # B. DELETE THEME REFERENCES
                        # This breaks the link to "Calibri (Headings)" or "Major Theme"
                        # If these exist, Word ignores the attributes we set above.
                        for attrib in ['w:asciiTheme', 'w:hAnsiTheme', 'w:eastAsiaTheme', 'w:cstheme']:
                            if qn(attrib) in rFonts.attrib:
                                del rFonts.attrib[qn(attrib)]
                                
                        changes.append(FormatChange(
                            type="font_family",
                            description=f"Changed font family",
                            before=old_font,
                            after=expected_font_family,
                        ))
            
            # --- NEW: Image, Caption and Source Formatting ---
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            # Check for image in this paragraph
            has_image = False
            if paragraph._element.find(qn('w:drawing')) is not None or \
               paragraph._element.find(qn('w:pict')) is not None:
                has_image = True
            if not has_image:
                for run in paragraph.runs:
                    if run._element.find(qn('w:drawing')) is not None or \
                       run._element.find(qn('w:pict')) is not None:
                        has_image = True
                        break
            
            # 1. Format Image Paragraph
            if has_image:
                if paragraph.alignment != WD_ALIGN_PARAGRAPH.CENTER:
                    old_align = str(paragraph.alignment)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    changes.append(FormatChange(
                        type="image_alignment",
                        description="Відцентровано зображення",
                        before=old_align,
                        after="center"
                    ))
            
            # 2. Format Caption Paragraph (Рис. X.X.)
            import re
            # Check for general caption-like start
            general_match = re.match(r'^(Рис\.|Зоб\.|Рисунок|Фото)\s*', paragraph.text.strip())
            if general_match:
                text = paragraph.text.strip()
                # Try to fix numbering format to exactly num.num
                # This pattern catches the first two numbers and ignores any subsequent ones (like .1.1)
                fixed_text = re.sub(
                    r'^(Рис|Зоб|Рисунок|Фото)\.?\s*(\d+)\s*[\.\,]\s*(\d+)(?:\.[\d\.]+)?\s*[\.\,]?\s*(.*)',
                    r'\1. \2.\3. \4',
                    text
                )
                
                if fixed_text != text:
                    paragraph.text = fixed_text
                    changes.append(FormatChange(
                        type="caption_format_fix",
                        description="Виправлено формат номера підпису",
                        before=text[:30],
                        after=fixed_text[:30]
                    ))

                if paragraph.alignment != WD_ALIGN_PARAGRAPH.CENTER:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    changes.append(FormatChange(
                        type="caption_alignment",
                        description=f"Відцентровано підпис",
                        before="unset",
                        after="center"
                    ))
            
            # 3. Format Source Paragraph (Джерело:)
            if paragraph.text.strip().lower().startswith("джерело:"):
                # Align center
                if paragraph.alignment != WD_ALIGN_PARAGRAPH.CENTER:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    changes.append(FormatChange(
                        type="source_alignment",
                        description="Відцентровано рядок джерела",
                        before="unset",
                        after="center"
                    ))
                # Set Italic
                for run in paragraph.runs:
                    if run.text.strip() and not run.font.italic:
                        run.font.italic = True
                        changes.append(FormatChange(
                            type="source_style",
                            description="Додано курсив до джерела",
                            before="normal",
                            after="italic"
                        ))
                        break # Only need to log once per paragraph
        
        # Apply margins
        if doc.sections:
            section = doc.sections[0]
            from docx.shared import Inches
            
            # Convert mm to inches (margins in params are in mm)
            def mm_to_inches(mm: float) -> float:
                return mm / 25.4
            
            # Top margin
            if params.margins.top:
                old_top = section.top_margin.inches if section.top_margin else 1.0
                new_top = mm_to_inches(params.margins.top)
                if abs(old_top - new_top) > self.MARGIN_TOLERANCE:
                    section.top_margin = Inches(new_top)
                    changes.append(FormatChange(
                        type="margin_top",
                        description="Updated top margin",
                        before=f"{old_top:.2f}\"",
                        after=f"{new_top:.2f}\"",
                    ))
            
            # Bottom margin
            if params.margins.bottom:
                old_bottom = section.bottom_margin.inches if section.bottom_margin else 1.0
                new_bottom = mm_to_inches(params.margins.bottom)
                if abs(old_bottom - new_bottom) > self.MARGIN_TOLERANCE:
                    section.bottom_margin = Inches(new_bottom)
                    changes.append(FormatChange(
                        type="margin_bottom",
                        description="Updated bottom margin",
                        before=f"{old_bottom:.2f}\"",
                        after=f"{new_bottom:.2f}\"",
                    ))
            
            # Left margin
            if params.margins.left:
                old_left = section.left_margin.inches if section.left_margin else 1.0
                new_left = mm_to_inches(params.margins.left)
                if abs(old_left - new_left) > self.MARGIN_TOLERANCE:
                    section.left_margin = Inches(new_left)
                    changes.append(FormatChange(
                        type="margin_left",
                        description="Updated left margin",
                        before=f"{old_left:.2f}\"",
                        after=f"{new_left:.2f}\"",
                    ))
            
            # Right margin
            if params.margins.right:
                old_right = section.right_margin.inches if section.right_margin else 1.0
                new_right = mm_to_inches(params.margins.right)
                if abs(old_right - new_right) > self.MARGIN_TOLERANCE:
                    section.right_margin = Inches(new_right)
                    changes.append(FormatChange(
                        type="margin_right",
                        description="Updated right margin",
                        before=f"{old_right:.2f}\"",
                        after=f"{new_right:.2f}\"",
                    ))
        
        # Apply page numbering
        if params.check_numbering and doc.sections:
            section = doc.sections[0]
            
            # Set up "Different first page" based on skip_first_page parameter
            try:
                sectPr = section._sectPr
                titlePg = sectPr.find(qn('w:titlePg'))
                
                if params.skip_first_page:
                    # Need different first page
                    if titlePg is None:
                        # Add titlePg element to enable "Different first page"
                        titlePg = etree.SubElement(sectPr, qn('w:titlePg'))
                        changes.append(FormatChange(
                            type="page_numbering",
                            description="Enabled 'Different first page'",
                            before="Same header/footer on all pages",
                            after="Different first page enabled",
                        ))
                else:
                    # Should NOT have different first page
                    if titlePg is not None:
                        # Remove titlePg element to disable "Different first page"
                        sectPr.remove(titlePg)
                        changes.append(FormatChange(
                            type="page_numbering",
                            description="Disabled 'Different first page'",
                            before="Different first page enabled",
                            after="Same header/footer on all pages",
                        ))
            except Exception as e:
                print(f"Error setting first page different: {e}")
            
            # Set page number start value
            try:
                sectPr = section._sectPr
                pgNumType = sectPr.find(qn('w:pgNumType'))
                
                if params.start_from_number != 1:
                    if pgNumType is None:
                        # Create pgNumType element
                        pgNumType = etree.SubElement(sectPr, qn('w:pgNumType'))
                    
                    old_start = pgNumType.get(qn('w:start'), '1')
                    pgNumType.set(qn('w:start'), str(params.start_from_number))
                    
                    if old_start != str(params.start_from_number):
                        changes.append(FormatChange(
                            type="page_numbering",
                            description="Set page number start",
                            before=f"Start: {old_start}",
                            after=f"Start: {params.start_from_number}",
                        ))
                elif pgNumType is not None:
                    # Remove custom start if setting to default (1)
                    start_attr = pgNumType.get(qn('w:start'))
                    if start_attr and start_attr != '1':
                        pgNumType.set(qn('w:start'), '1')
                        changes.append(FormatChange(
                            type="page_numbering",
                            description="Reset page number start to default",
                            before=f"Start: {start_attr}",
                            after="Start: 1",
                        ))
            except Exception as e:
                print(f"Error setting page number start: {e}")
            
            # Add page number field to footer if not present
            try:
                footer = section.footer
                footer_xml = footer._element
                
                # Check if footer already has page number
                has_page_num = False
                for instr_text in footer_xml.iter(qn('w:instrText')):
                    if instr_text.text and 'PAGE' in instr_text.text:
                        has_page_num = True
                        break
                
                if not has_page_num:
                    # Clear existing footer content
                    for p in footer.paragraphs:
                        p._element.getparent().remove(p._element)
                    
                    # Add new centered paragraph with page number
                    p = footer.add_paragraph()
                    p.alignment = 1  # Center alignment (WD_ALIGN_PARAGRAPH.CENTER)
                    
                    # Add page number field
                    run = p.add_run()
                    r = run._element
                    
                    # Create field begin
                    fldChar_begin = etree.SubElement(r, qn('w:fldChar'))
                    fldChar_begin.set(qn('w:fldCharType'), 'begin')
                    
                    # Create instruction text with PAGE field
                    instrText = etree.SubElement(r, qn('w:instrText'))
                    instrText.set(qn('xml:space'), 'preserve')
                    instrText.text = ' PAGE '
                    
                    # Create field separator
                    fldChar_sep = etree.SubElement(r, qn('w:fldChar'))
                    fldChar_sep.set(qn('w:fldCharType'), 'separate')
                    
                    # Create field result (placeholder text)
                    t = etree.SubElement(r, qn('w:t'))
                    t.text = '1'
                    
                    # Create field end
                    fldChar_end = etree.SubElement(r, qn('w:fldChar'))
                    fldChar_end.set(qn('w:fldCharType'), 'end')
                    
                    changes.append(FormatChange(
                        type="page_numbering",
                        description="Added page numbers to footer",
                        before="No page numbers",
                        after="Page numbers in footer (centered)",
                    ))
            except Exception as e:
                print(f"Error adding page numbers: {e}")
        
        # Save modified document to bytes
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        modified_content = output.getvalue()
        
        processing_time = int((time.time() - start_time) * 1000)
        
        result = FormatterResult(
            success=True,
            changes_applied=len(changes),
            changes=changes,
            processing_time_ms=processing_time,
            document_title=doc_title,
        )
        
        return modified_content, result


# Dependency for FastAPI
LocalDocumentServiceDependency = Annotated[LocalDocumentService, Depends(LocalDocumentService)]
