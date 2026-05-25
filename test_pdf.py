import sys
import os
from docx import Document
from io import BytesIO

# Add to path to import core
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from core.pdf_utils import get_page_start_text_via_pdf

def test():
    doc = Document()
    doc.add_paragraph("This is the first page.")
    doc.add_page_break()
    doc.add_paragraph("This is the second page.")
    
    bio = BytesIO()
    doc.save(bio)
    
    bytes_data = bio.getvalue()
    
    text = get_page_start_text_via_pdf(bytes_data, 1) # Page index 1 (second page)
    print(f"Extracted text from second page: '{text}'")

if __name__ == '__main__':
    test()
